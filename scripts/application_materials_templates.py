from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pydantic import BaseModel, Field


@dataclass(frozen=True)
class TemplateConfig:
    role_family: str
    resume_template_path: str
    cover_letter_template_path: str
    preferred_resume_source_file: str
    skill_line_count: int
    project_bullet_count: int
    senior_experience_bullet_count: int
    cover_letter_body_count: int
    summary_prefixes: tuple[str, ...]
    professional_development_body_source: str
    project_bullet_sources: tuple[str, ...]
    senior_experience_bullet_sources: tuple[str, ...]
    skills_heading: str
    projects_heading: str
    professional_experience_heading: str
    senior_role_heading: str
    next_role_heading: str


class ResumeContent(BaseModel):
    summary: str
    skills_lines: list[str] = Field(min_length=1, max_length=8)
    professional_development_body: str
    project_bullets: list[str] = Field(min_length=1, max_length=8)
    senior_experience_bullets: list[str] = Field(min_length=1, max_length=8)


class CoverLetterContent(BaseModel):
    body_paragraphs: list[str] = Field(min_length=1, max_length=6)


SOFTWARE_ENGINEERING_TEMPLATE = TemplateConfig(
    role_family="software_engineering",
    resume_template_path="Archive/vargas_software_engineer_resume.docx",
    cover_letter_template_path="Archive/vargas_software_engineer_cover_letter.docx",
    preferred_resume_source_file="resumes_cov_letters/archive/vargas_senior_software_engineer_resume.docx",
    skill_line_count=4,
    project_bullet_count=4,
    senior_experience_bullet_count=6,
    cover_letter_body_count=5,
    summary_prefixes=("Software engineer with", "Technical support engineer with"),
    professional_development_body_source="education/paragraph_47",
    project_bullet_sources=(
        "projects/paragraph_20",
        "projects/paragraph_21",
        "projects/paragraph_22",
        "projects/paragraph_23",
    ),
    senior_experience_bullet_sources=(
        "experience/paragraph_28",
        "experience/paragraph_29",
        "experience/paragraph_30",
        "experience/paragraph_31",
        "experience/paragraph_32",
        "experience/paragraph_33",
    ),
    skills_heading="TECHNICAL SKILLS",
    projects_heading="RECENT PROJECTS",
    professional_experience_heading="PROFESSIONAL EXPERIENCE",
    senior_role_heading="IQVIA | Senior Software Engineer",
    next_role_heading="IQVIA | Software Engineer III",
)


CUSTOMER_FACING_TECHNICAL_TEMPLATE = TemplateConfig(
    role_family="customer_facing_technical",
    resume_template_path="reference_examples/customer_facing_technical/vargas_technical_implementation_specialist_resume.docx",
    cover_letter_template_path="reference_examples/customer_facing_technical/vargas_technical_implementation_specialist_cover_letter.docx",
    preferred_resume_source_file="resumes_cov_letters/archive/vargas_technical_support_engineer_resume.docx",
    skill_line_count=6,
    project_bullet_count=4,
    senior_experience_bullet_count=6,
    cover_letter_body_count=3,
    summary_prefixes=(
        "Software engineer with 6+ years experience supporting teams in healthcare and life-sciences.",
        "6 years of software engineering and technical support experience",
        "Technical support engineer with",
        "Software engineer with",
    ),
    professional_development_body_source="paragraph_21",
    project_bullet_sources=(
        "projects/paragraph_25",
        "projects/paragraph_26",
        "projects/paragraph_27",
        "projects/paragraph_28",
    ),
    senior_experience_bullet_sources=(
        "experience/paragraph_33",
        "experience/paragraph_34",
        "experience/paragraph_35",
        "experience/paragraph_36",
        "experience/paragraph_37",
    ),
    skills_heading="TECHNICAL SKILLS",
    projects_heading="RECENT PROJECTS",
    professional_experience_heading="PROFESSIONAL EXPERIENCE",
    senior_role_heading="IQVIA | Senior Software Engineer",
    next_role_heading="IQVIA | Software Engineer III",
)


TECHNICAL_WRITING_TEMPLATE = TemplateConfig(
    role_family="technical_writing",
    resume_template_path="reference_examples/technical_writing/vargas_technical_writer_resume_core_competencies.docx",
    cover_letter_template_path="reference_examples/technical_writing/vargas_technical_writer_cover_letter_sprypoint.docx",
    preferred_resume_source_file="resumes_cov_letters/archive/vargas_technical_support_engineer_resume.docx",
    skill_line_count=5,
    project_bullet_count=1,
    senior_experience_bullet_count=5,
    cover_letter_body_count=4,
    summary_prefixes=(
        "Technical writer and former software engineer",
        "Technical writer and software engineer",
    ),
    professional_development_body_source="paragraph_21",
    project_bullet_sources=("projects/paragraph_25",),
    senior_experience_bullet_sources=(
        "experience/paragraph_33",
        "experience/paragraph_34",
        "experience/paragraph_35",
        "experience/paragraph_36",
        "experience/paragraph_37",
    ),
    skills_heading="CORE COMPETENCIES",
    projects_heading="PROJECTS",
    professional_experience_heading="PROFESSIONAL EXPERIENCE",
    senior_role_heading="IQVIA | Senior Software Engineer",
    next_role_heading="IQVIA | Software Engineer III",
)


CUSTOMER_FACING_ROLE_FAMILIES = {
    "solutions_engineering",
    "technical_support_engineering",
}


def normalize_match_text(text: str) -> str:
    import re

    lowered = text.lower()
    lowered = re.sub(r"[_/|]+", " ", lowered)
    lowered = re.sub(r"[^a-z0-9+&.-]+", " ", lowered)
    return " ".join(lowered.split())


def infer_generation_role_family(title: str, description: str = "") -> str:
    normalized = normalize_match_text(f"{title} {description}")
    if any(token in normalized for token in ("technical writer", "documentation writer", "documentation specialist")):
        return "technical_writing"
    if any(
        token in normalized
        for token in (
            "solutions engineer",
            "customer solutions engineer",
            "solutions architect",
            "technical account manager",
            "customer success engineer",
            "implementation specialist",
            "integration specialist",
            "technical trainer",
            "customer onboarding specialist",
        )
    ):
        return "solutions_engineering"
    if any(
        token in normalized
        for token in (
            "technical support engineer",
            "technical support specialist",
            "product support specialist",
            "software support engineer",
            "application specialist",
            "tier 2 support",
            "tier 3 support",
            "scientific support specialist",
            "support engineer",
        )
    ):
        return "technical_support_engineering"
    return "software_engineering"


def template_config_for_role_family(role_family: str) -> TemplateConfig:
    if role_family == "technical_writing":
        return TECHNICAL_WRITING_TEMPLATE
    if role_family in CUSTOMER_FACING_ROLE_FAMILIES:
        return CUSTOMER_FACING_TECHNICAL_TEMPLATE
    return SOFTWARE_ENGINEERING_TEMPLATE


def clean_display_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def find_paragraph_index(document: Document, predicate) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if predicate(clean_display_text(paragraph.text)):
            return index
    raise ValueError("Could not locate required template paragraph")


def non_empty_indexes(document: Document, start: int, end: int) -> list[int]:
    indexes: list[int] = []
    for idx in range(start, end):
        if clean_display_text(document.paragraphs[idx].text):
            indexes.append(idx)
    return indexes


def require_minimum_indexes(indexes: list[int], minimum: int, section_name: str) -> list[int]:
    if len(indexes) < minimum:
        raise ValueError(
            f"Template is missing expected paragraphs for {section_name}: expected at least {minimum}, found {len(indexes)}"
        )
    return indexes


def set_paragraph_text(document: Document, index: int, text: str) -> None:
    document.paragraphs[index].text = text


def validate_resume_content_for_template(content: ResumeContent, template_config: TemplateConfig) -> None:
    if len(content.skills_lines) != template_config.skill_line_count:
        raise ValueError(
            f"Resume content has {len(content.skills_lines)} skill lines; expected {template_config.skill_line_count} for {template_config.resume_template_path}"
        )
    if len(content.project_bullets) != template_config.project_bullet_count:
        raise ValueError(
            f"Resume content has {len(content.project_bullets)} project bullets; expected {template_config.project_bullet_count} for {template_config.resume_template_path}"
        )
    if len(content.senior_experience_bullets) != template_config.senior_experience_bullet_count:
        raise ValueError(
            f"Resume content has {len(content.senior_experience_bullets)} senior bullets; expected {template_config.senior_experience_bullet_count} for {template_config.resume_template_path}"
        )


def extract_template_resume_content(application_materials_dir: Path, template_config: TemplateConfig) -> tuple[list[str], str, list[str], list[str]]:
    document = Document(application_materials_dir / template_config.resume_template_path)

    skills_heading_index = find_paragraph_index(document, lambda text: text == template_config.skills_heading)
    prof_dev_heading_index = find_paragraph_index(document, lambda text: text == "RECENT PROFESSIONAL DEVELOPMENT")
    projects_heading_index = find_paragraph_index(document, lambda text: text == template_config.projects_heading)
    experience_heading_index = find_paragraph_index(document, lambda text: text == template_config.professional_experience_heading)
    senior_role_index = find_paragraph_index(document, lambda text: text == template_config.senior_role_heading)
    next_role_index = find_paragraph_index(document, lambda text: text == template_config.next_role_heading)

    skill_indexes = require_minimum_indexes(
        non_empty_indexes(document, skills_heading_index + 1, prof_dev_heading_index),
        template_config.skill_line_count,
        "resume technical skills",
    )
    prof_dev_indexes = require_minimum_indexes(
        non_empty_indexes(document, prof_dev_heading_index + 1, projects_heading_index),
        3,
        "recent professional development",
    )
    project_indexes = require_minimum_indexes(
        non_empty_indexes(document, projects_heading_index + 1, experience_heading_index),
        template_config.project_bullet_count + 1,
        "recent projects",
    )
    senior_indexes = require_minimum_indexes(
        non_empty_indexes(document, senior_role_index + 1, next_role_index),
        template_config.senior_experience_bullet_count + 1,
        "senior software engineer experience",
    )

    skills_lines = [clean_display_text(document.paragraphs[idx].text) for idx in skill_indexes[: template_config.skill_line_count]]
    professional_development_body = clean_display_text(document.paragraphs[prof_dev_indexes[2]].text)
    project_bullets = [
        clean_display_text(document.paragraphs[idx].text)
        for idx in project_indexes[1 : 1 + template_config.project_bullet_count]
    ]
    senior_experience_bullets = [
        clean_display_text(document.paragraphs[idx].text)
        for idx in senior_indexes[1 : 1 + template_config.senior_experience_bullet_count]
    ]
    return skills_lines, professional_development_body, project_bullets, senior_experience_bullets


def render_resume_docx(application_materials_dir: Path, output_path: Path, content: ResumeContent, template_config: TemplateConfig) -> Document:
    validate_resume_content_for_template(content, template_config)
    document = Document(application_materials_dir / template_config.resume_template_path)

    summary_index = find_paragraph_index(
        document,
        lambda text: any(text.startswith(prefix) for prefix in template_config.summary_prefixes),
    )
    skills_heading_index = find_paragraph_index(document, lambda text: text == template_config.skills_heading)
    prof_dev_heading_index = find_paragraph_index(document, lambda text: text == "RECENT PROFESSIONAL DEVELOPMENT")
    projects_heading_index = find_paragraph_index(document, lambda text: text == template_config.projects_heading)
    experience_heading_index = find_paragraph_index(document, lambda text: text == template_config.professional_experience_heading)
    senior_role_index = find_paragraph_index(document, lambda text: text == template_config.senior_role_heading)
    next_role_index = find_paragraph_index(document, lambda text: text == template_config.next_role_heading)

    skill_indexes = require_minimum_indexes(
        non_empty_indexes(document, skills_heading_index + 1, prof_dev_heading_index),
        template_config.skill_line_count,
        "resume technical skills",
    )
    prof_dev_indexes = require_minimum_indexes(
        non_empty_indexes(document, prof_dev_heading_index + 1, projects_heading_index),
        3,
        "recent professional development",
    )
    project_indexes = require_minimum_indexes(
        non_empty_indexes(document, projects_heading_index + 1, experience_heading_index),
        template_config.project_bullet_count + 1,
        "recent projects",
    )
    senior_indexes = require_minimum_indexes(
        non_empty_indexes(document, senior_role_index + 1, next_role_index),
        template_config.senior_experience_bullet_count + 1,
        "senior software engineer experience",
    )

    set_paragraph_text(document, summary_index, content.summary)

    for idx, line in zip(skill_indexes[: template_config.skill_line_count], content.skills_lines):
        set_paragraph_text(document, idx, line)

    set_paragraph_text(document, prof_dev_indexes[2], content.professional_development_body)

    for idx, bullet in zip(
        project_indexes[1 : 1 + template_config.project_bullet_count],
        content.project_bullets,
    ):
        set_paragraph_text(document, idx, bullet)

    for idx, bullet in zip(
        senior_indexes[1 : 1 + template_config.senior_experience_bullet_count],
        content.senior_experience_bullets,
    ):
        set_paragraph_text(document, idx, bullet)

    document.save(output_path)
    return document


def render_cover_letter_docx(application_materials_dir: Path, output_path: Path, content: CoverLetterContent, template_config: TemplateConfig) -> Document:
    if len(content.body_paragraphs) != template_config.cover_letter_body_count:
        raise ValueError(
            f"Cover-letter content has {len(content.body_paragraphs)} body paragraphs; expected {template_config.cover_letter_body_count} for {template_config.cover_letter_template_path}"
        )
    document = Document(application_materials_dir / template_config.cover_letter_template_path)
    salutation_index = find_paragraph_index(document, lambda text: text.startswith("Dear "))
    closing_index = find_paragraph_index(document, lambda text: text == "Sincerely,")
    body_indexes = require_minimum_indexes(
        non_empty_indexes(document, salutation_index + 1, closing_index),
        template_config.cover_letter_body_count,
        "cover letter body",
    )

    for idx, paragraph_text in zip(body_indexes, content.body_paragraphs):
        set_paragraph_text(document, idx, paragraph_text)

    document.save(output_path)
    return document


def document_to_text(document: Document) -> str:
    lines = [clean_display_text(paragraph.text) for paragraph in document.paragraphs]
    return "\n".join(line for line in lines if line)